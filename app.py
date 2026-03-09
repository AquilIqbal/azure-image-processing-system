import io
import mimetypes
import os
import uuid
import zipfile
from datetime import datetime

from azure.core.exceptions import AzureError
from docx import Document
from flask import Flask, Response, abort, flash, redirect, render_template, request, send_file, url_for
from flask_login import LoginManager, current_user, login_required, login_user, logout_user
from PIL import Image
from sqlalchemy import func
from sqlalchemy.exc import SQLAlchemyError
from werkzeug.security import check_password_hash, generate_password_hash
from werkzeug.exceptions import RequestEntityTooLarge

from azure_blob_helper import AzureBlobHelper
from database import ImageMetadata, User, db, init_db

from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY", "replace-this-with-a-secure-key")
app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024

# Azure SQL connection string should be configured in an environment variable.
app.config["SQLALCHEMY_DATABASE_URI"] = os.getenv("DATABASE_URL")
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = "login"

init_db(app)
blob_helper = AzureBlobHelper()

IMAGE_EXTENSIONS = {"jpg", "jpeg", "png", "webp"}
DOCUMENT_EXTENSIONS = {"pdf", "docx", "txt"}
ALLOWED_EXTENSIONS = IMAGE_EXTENSIONS | DOCUMENT_EXTENSIONS
GALLERY_IMAGE_FORMATS = {"jpg", "jpeg", "png", "gif", "bmp", "webp", "tiff"}
GALLERY_DOCUMENT_FORMATS = {"pdf", "docx", "txt", "zip"}

RESIZE_PRESETS = {
    "none": None,
    "512": (512, 512),
    "1024": (1024, 1024),
    "2048": (2048, 2048),
}
COMPRESSION_LEVELS = {
    "none": None,
    "low": {"quality": 90, "compress_level": 3},
    "medium": {"quality": 75, "compress_level": 6},
    "high": {"quality": 60, "compress_level": 9},
}


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


def _extension(filename):
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def _safe_base_name(filename):
    return filename.rsplit(".", 1)[0] if "." in filename else filename


def _build_processing_summary(item):
    original_ext = _extension(item.original_filename)
    final_ext = item.file_format.lower()

    is_document = original_ext in DOCUMENT_EXTENSIONS or final_ext == "zip"
    item.category_label = "Document" if is_document else "Image"
    item.resize_dimensions = f"{item.width}x{item.height}" if item.width and item.height else "Original size"
    item.compression_badge = (item.compression_level or "none").upper()

    ops = []
    if is_document:
        if original_ext == "docx" and final_ext == "txt":
            ops.append("Converted DOCX to TXT")
        elif original_ext == "txt" and final_ext == "docx":
            ops.append("Converted TXT to DOCX")
        if final_ext == "zip":
            ops.append("Compressed to ZIP")
    else:
        if item.resize_applied and item.width and item.height:
            ops.append(f"Resize {item.width}x{item.height}")
        if item.compress_applied:
            ops.append("Compression applied")
        if original_ext and final_ext and original_ext != final_ext:
            ops.append(f"Converted to {final_ext.upper()}")

    if not ops:
        ops.append("No processing")

    item.operations_display = ", ".join(ops)

    original_size = item.original_size_kb if item.original_size_kb is not None else round(item.file_size_bytes / 1024, 2)
    processed_size = item.processed_size_kb if item.processed_size_kb is not None else round(item.file_size_bytes / 1024, 2)

    item.original_size = f"{original_size:.2f}"
    if is_document and item.operations_display == "No processing":
        item.processed_size = "No processing"
        item.savings_display = None
    else:
        item.processed_size = f"{processed_size:.2f}"
        if not is_document and original_size > 0 and processed_size < original_size:
            saved = original_size - processed_size
            percentage = (saved / original_size) * 100
            item.savings_display = f"Saved: {saved:.2f} KB ({percentage:.0f}%)"
        else:
            item.savings_display = None


def process_image(file_storage, original_extension, resize_option, custom_width, custom_height, compression_option, format_option):
    """Apply resize/compression/format conversion to image files."""
    image = Image.open(file_storage.stream)

    target_size = RESIZE_PRESETS.get(resize_option)
    if resize_option == "custom" and custom_width and custom_height:
        target_size = (custom_width, custom_height)
    if target_size:
        image = image.resize(target_size, Image.Resampling.LANCZOS)

    requested_format = (format_option or "keep").lower().strip()
    selected_extension = original_extension if requested_format == "keep" else requested_format
    if selected_extension == "jpeg":
        selected_extension = "jpg"
    if selected_extension not in {"jpg", "png", "webp", "gif", "bmp", "tiff"}:
        selected_extension = "png"

    selected_format = "JPEG" if selected_extension == "jpg" else selected_extension.upper()

    if selected_extension in {"jpg", "webp"} and image.mode in {"RGBA", "LA", "P"}:
        background = Image.new("RGB", image.size, (255, 255, 255))
        if image.mode != "RGBA":
            image = image.convert("RGBA")
        background.paste(image, mask=image.split()[-1])
        image = background

    compression = COMPRESSION_LEVELS.get(compression_option, COMPRESSION_LEVELS["none"])
    save_kwargs = {"optimize": True}
    if selected_extension == "png":
        save_kwargs["compress_level"] = compression["compress_level"] if compression else 1
    elif selected_extension in {"jpg", "webp"}:
        save_kwargs["quality"] = compression["quality"] if compression else 95

    buffer = io.BytesIO()
    image.save(buffer, format=selected_format, **save_kwargs)
    buffer.seek(0)

    resize_applied = bool(target_size)
    compress_applied = compression_option != "none"
    width = target_size[0] if target_size else None
    height = target_size[1] if target_size else None

    return buffer, selected_extension, resize_applied, compress_applied, width, height


def process_document(file_storage, original_filename, original_extension, conversion_option, compress_document):
    """Apply optional DOCX/TXT conversion and optional ZIP compression for documents."""
    raw_bytes = file_storage.stream.read()
    current_bytes = raw_bytes
    current_extension = original_extension

    conversion_option = (conversion_option or "keep").lower().strip()

    if conversion_option == "txt" and original_extension == "docx":
        doc = Document(io.BytesIO(raw_bytes))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        current_bytes = text.encode("utf-8")
        current_extension = "txt"
    elif conversion_option == "docx" and original_extension == "txt":
        text = raw_bytes.decode("utf-8", errors="ignore")
        doc = Document()
        doc.add_paragraph(text)
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        current_bytes = doc_buffer.read()
        current_extension = "docx"

    if compress_document:
        zipped = io.BytesIO()
        archive_name = f"{_safe_base_name(original_filename)}.{current_extension}"
        with zipfile.ZipFile(zipped, mode="w", compression=zipfile.ZIP_DEFLATED) as archive:
            archive.writestr(archive_name, current_bytes)
        zipped.seek(0)
        return zipped, "zip", True

    stream = io.BytesIO(current_bytes)
    stream.seek(0)
    return stream, current_extension, False


def _upload_and_store_metadata(file_storage, file_kind):
    original_extension = _extension(file_storage.filename)

    file_storage.stream.seek(0, io.SEEK_END)
    original_size_bytes = file_storage.stream.tell()
    file_storage.stream.seek(0)
    original_size_kb = round(original_size_bytes / 1024, 2)

    resize_applied = False
    compress_applied = False
    compression_level = "none"
    width = None
    height = None

    if file_kind == "image":
        resize_option = request.form.get("resize_option", "none")
        compression_option = request.form.get("compression_option", "medium")
        format_option = request.form.get("format_option", "keep")
        custom_width = request.form.get("custom_width", type=int)
        custom_height = request.form.get("custom_height", type=int)

        processed_stream, extension, resize_applied, compress_applied, width, height = process_image(
            file_storage,
            original_extension,
            resize_option,
            custom_width,
            custom_height,
            compression_option,
            format_option,
        )
        compression_level = compression_option
        blob_name = f"processed/{uuid.uuid4().hex}.{extension}"
    else:
        conversion_option = request.form.get("document_conversion", "keep")
        compress_document = request.form.get("document_compress") == "on"
        processed_stream, extension, compress_applied = process_document(
            file_storage,
            file_storage.filename,
            original_extension,
            conversion_option,
            compress_document,
        )
        compression_level = "zip" if compress_document else "none"
        blob_name = f"documents/{uuid.uuid4().hex}.{extension}"

    processed_size_kb = round(processed_stream.getbuffer().nbytes / 1024, 2)
    try:
        blob_url = blob_helper.upload_image(blob_name, processed_stream)
    except AzureError as exc:
        raise RuntimeError("BLOB_ERROR") from exc

    metadata = ImageMetadata(
        original_filename=file_storage.filename,
        blob_name=blob_name,
        blob_url=blob_url,
        container_name=blob_helper.container_name,
        user_id=current_user.id if current_user.is_authenticated else None,
        file_format=extension,
        file_size_bytes=processed_stream.getbuffer().nbytes,
        original_size_kb=original_size_kb,
        processed_size_kb=processed_size_kb,
        resize_applied=resize_applied,
        compress_applied=compress_applied,
        compression_level=compression_level,
        download_count=0,
        width=width,
        height=height,
        uploaded_at=datetime.utcnow(),
    )
    try:
        db.session.add(metadata)
        db.session.commit()
    except SQLAlchemyError as exc:
        db.session.rollback()
        raise RuntimeError("DB_ERROR") from exc


def _handle_upload_request(expected_kind):
    file = request.files.get("image") or request.files.get("document") or request.files.get("file")
    if not file or file.filename == "":
        flash("Please select a file to upload.", "danger")
        return redirect(url_for("images" if expected_kind == "image" else "documents"))

    extension = _extension(file.filename)
    if extension not in ALLOWED_EXTENSIONS:
        flash("Unsupported file type", "danger")
        return redirect(url_for("images" if expected_kind == "image" else "documents"))

    actual_kind = "image" if extension in IMAGE_EXTENSIONS else "document"
    if expected_kind == "image" and actual_kind != "image":
        flash("Please upload an image file on the Images page.", "warning")
        return redirect(url_for("images"))
    if expected_kind == "document" and actual_kind != "document":
        flash("Please upload a document file on the Documents page.", "warning")
        return redirect(url_for("documents"))

    try:
        _upload_and_store_metadata(file, actual_kind)
        flash("File uploaded successfully.", "success")
    except RuntimeError as exc:
        if str(exc) == "BLOB_ERROR":
            flash("Blob storage error occurred while uploading file.", "danger")
        elif str(exc) == "DB_ERROR":
            flash("Database error occurred while saving metadata.", "danger")
        else:
            flash("Upload failure. Please try again.", "danger")
    except Exception:
        db.session.rollback()
        flash("Upload failure. Please try again.", "danger")

    return redirect(url_for("gallery"))


@app.errorhandler(RequestEntityTooLarge)
def handle_file_too_large(e):
    flash("File too large. Maximum allowed size is 10MB.", "danger")
    return redirect(url_for("index"))


@app.route("/register", methods=["GET", "POST"])
def register():
    if current_user.is_authenticated:
        return redirect(url_for("gallery"))

    if request.method == "POST":
        username = request.form.get("username", "").strip()
        email = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "")

        if not username or not email or not password:
            flash("All fields are required.", "danger")
            return redirect(url_for("register"))

        if User.query.filter((User.email == email) | (User.username == username)).first():
            flash("Username or email already exists.", "danger")
            return redirect(url_for("register"))

        user = User(
            username=username,
            email=email,
            password_hash=generate_password_hash(password),
        )
        try:
            db.session.add(user)
            db.session.commit()
            flash("Registration successful. Please login.", "success")
            return redirect(url_for("login"))
        except SQLAlchemyError:
            db.session.rollback()
            flash("Database error occurred while creating account.", "danger")
            return redirect(url_for("register"))

    return render_template("register.html", active_page="register")


@app.route("/login", methods=["GET", "POST"])
def login():
    if current_user.is_authenticated:
        return redirect(url_for("gallery"))

    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "")
        user = User.query.filter_by(email=email).first()

        if not user or not check_password_hash(user.password_hash, password):
            flash("Invalid email or password.", "danger")
            return redirect(url_for("login"))

        login_user(user)
        flash("Logged in successfully.", "success")
        next_url = request.args.get("next")
        return redirect(next_url or url_for("gallery"))

    return render_template("login.html", active_page="login")


@app.route("/logout")
@login_required
def logout():
    logout_user()
    flash("Logged out successfully.", "success")
    return redirect(url_for("login"))


@app.route("/")
def index():
    return render_template("index.html", active_page="home")


@app.route("/images")
@login_required
def images():
    return render_template("images.html", active_page="images")


@app.route("/documents")
@login_required
def documents():
    return render_template("documents.html", active_page="documents")


@app.route("/upload", methods=["POST", "GET"])
@login_required
def upload_file():
    if request.method == "GET":
        return redirect(url_for("index"))
    file = request.files.get("image") or request.files.get("document") or request.files.get("file")
    if not file or file.filename == "":
        flash("Please select a file to upload.", "danger")
        return redirect(url_for("index"))

    extension = _extension(file.filename)
    if extension in IMAGE_EXTENSIONS:
        return _handle_upload_request("image")
    if extension in DOCUMENT_EXTENSIONS:
        return _handle_upload_request("document")

    flash("Unsupported file type", "danger")
    return redirect(url_for("index"))


@app.route("/upload/image", methods=["POST", "GET"])
@login_required
def upload_image_only():
    if request.method == "GET":
        return redirect(url_for("images"))
    return _handle_upload_request("image")


@app.route("/upload/document", methods=["POST", "GET"])
@login_required
def upload_document_only():
    if request.method == "GET":
        return redirect(url_for("documents"))
    return _handle_upload_request("document")


@app.route("/gallery")
@login_required
def gallery():
    search = request.args.get("search", "").strip()
    type_filter = request.args.get("type", "all").strip().lower()
    sort = request.args.get("sort", "newest").strip().lower()

    query = ImageMetadata.query.filter_by(user_id=current_user.id)
    if search:
        query = query.filter(ImageMetadata.original_filename.ilike(f"%{search}%"))
    if type_filter == "image":
        query = query.filter(ImageMetadata.file_format.in_(GALLERY_IMAGE_FORMATS))
    elif type_filter == "document":
        query = query.filter(ImageMetadata.file_format.in_(GALLERY_DOCUMENT_FORMATS))

    if sort == "oldest":
        query = query.order_by(ImageMetadata.uploaded_at.asc())
    elif sort == "largest":
        query = query.order_by(ImageMetadata.file_size_bytes.desc())
    elif sort == "downloads":
        query = query.order_by(ImageMetadata.download_count.desc(), ImageMetadata.uploaded_at.desc())
    else:
        sort = "newest"
        query = query.order_by(ImageMetadata.uploaded_at.desc())

    files = query.all()
    for item in files:
        _build_processing_summary(item)
    return render_template(
        "gallery.html",
        images=files,
        active_page="gallery",
        search=search,
        type_filter=type_filter,
        sort=sort,
    )


@app.route("/download/<int:image_id>")
@login_required
def download_image(image_id):
    image_meta = ImageMetadata.query.get_or_404(image_id)
    if image_meta.user_id != current_user.id:
        abort(403)
    try:
        stream = blob_helper.download_image(image_meta.blob_name)
        image_meta.download_count = (image_meta.download_count or 0) + 1
        try:
            db.session.commit()
        except SQLAlchemyError:
            db.session.rollback()
        mime_type = mimetypes.types_map.get(f".{image_meta.file_format.lower()}", "application/octet-stream")
        download_name = f"processed_{_safe_base_name(image_meta.original_filename)}.{image_meta.file_format}"
        return send_file(stream, mimetype=mime_type, as_attachment=True, download_name=download_name)
    except AzureError:
        flash("Blob storage error occurred while downloading file.", "danger")
        return redirect(url_for("gallery"))
    except Exception as exc:
        flash(f"Download failed: {exc}", "danger")
        return redirect(url_for("gallery"))


@app.route("/preview/<int:file_id>")
@login_required
def preview_file(file_id):
    file_meta = ImageMetadata.query.get_or_404(file_id)
    if file_meta.user_id != current_user.id:
        abort(403)
    extension = file_meta.file_format.lower()

    preview_mime_types = {
        "pdf": "application/pdf",
        "txt": "text/plain; charset=utf-8",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    mime_type = preview_mime_types.get(extension)
    if not mime_type:
        return Response("Preview not supported for this file type.", status=415, mimetype="text/plain")

    try:
        stream = blob_helper.download_image(file_meta.blob_name)
        file_data = stream.getvalue()
        return Response(
            file_data,
            mimetype=mime_type,
            headers={"Content-Disposition": "inline"},
        )
    except AzureError:
        return Response("Blob storage error during preview.", status=502, mimetype="text/plain")
    except Exception:
        return Response("Unable to preview this file.", status=500, mimetype="text/plain")


@app.route("/delete/<int:image_id>", methods=["POST"])
@login_required
def delete_image(image_id):
    image_meta = ImageMetadata.query.get_or_404(image_id)
    if image_meta.user_id != current_user.id:
        abort(403)
    try:
        blob_helper.delete_image(image_meta.blob_name)
        db.session.delete(image_meta)
        db.session.commit()
        flash("File deleted successfully.", "success")
    except AzureError:
        db.session.rollback()
        flash("Blob storage error occurred while deleting file.", "danger")
    except SQLAlchemyError:
        db.session.rollback()
        flash("Database error occurred while deleting metadata.", "danger")
    except Exception as exc:
        db.session.rollback()
        flash(f"Delete failed: {exc}", "danger")

    return redirect(url_for("gallery"))


@app.route("/analytics")
@login_required
def analytics():
    total_uploads = db.session.query(func.count(ImageMetadata.id)).filter(
        ImageMetadata.user_id == current_user.id
    ).scalar() or 0
    total_images = db.session.query(func.count(ImageMetadata.id)).filter(
        ImageMetadata.user_id == current_user.id,
        ImageMetadata.file_format.in_(GALLERY_IMAGE_FORMATS)
    ).scalar() or 0
    total_documents = db.session.query(func.count(ImageMetadata.id)).filter(
        ImageMetadata.user_id == current_user.id,
        ImageMetadata.file_format.in_(GALLERY_DOCUMENT_FORMATS)
    ).scalar() or 0
    total_storage_bytes = db.session.query(func.coalesce(func.sum(ImageMetadata.file_size_bytes), 0)).filter(
        ImageMetadata.user_id == current_user.id
    ).scalar() or 0
    avg_file_size = db.session.query(func.coalesce(func.avg(ImageMetadata.file_size_bytes), 0)).filter(
        ImageMetadata.user_id == current_user.id
    ).scalar() or 0

    stats = {
        "total_uploads": total_uploads,
        "total_images": total_images,
        "total_documents": total_documents,
        "total_storage_mb": round(total_storage_bytes / (1024 * 1024), 2),
        "average_file_kb": round(avg_file_size / 1024, 2),
    }
    return render_template("analytics.html", active_page="analytics", stats=stats)


@app.route("/history")
@login_required
def history():
    files = ImageMetadata.query.filter_by(user_id=current_user.id).order_by(ImageMetadata.uploaded_at.desc()).all()
    return render_template("history.html", active_page="history", files=files)

if __name__ == "__main__":
    with app.app_context():
        db.create_all()
    app.run(debug=True)
